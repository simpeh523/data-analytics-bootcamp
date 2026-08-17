"""Builds outputs/Edwards_Caleb_Capstone_Report.docx — the Stage 6 comprehensive report.

Every figure in this document is quoted from outputs/FINDINGS.md, which is itself
produced by outputs/capstone_analysis.ipynb against outputs/part_d_co_clean.csv.
No number is introduced here that does not already exist in FINDINGS.md or CLEANING.md.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = "/sessions/nice-confident-volta/mnt/DAB Capstone--outputs"
CH = os.path.join(OUT, "charts")

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GREY = RGBColor(0x4A, 0x55, 0x68)
NAVY_HEX = "1A365D"
HDRFILL = "E7ECF4"
RULE = "B4C0D4"

doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
s.top_margin = s.bottom_margin = Inches(0.85)
s.left_margin = s.right_margin = Inches(1.0)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.10


# ---------- low-level helpers ----------

def shade(cell, hexfill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def _fixed_layout():
    el = OxmlElement("w:tblLayout")
    el.set(qn("w:type"), "fixed")
    return el


def set_borders(table, colour=RULE):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4" if edge in ("insideH", "insideV") else "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), colour)
        borders.append(el)
    tblPr.append(borders)


def left_bar(par, colour=NAVY_HEX):
    """Vertical accent rule down the left of a paragraph (used for the question callout)."""
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    el = OxmlElement("w:left")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "18")
    el.set(qn("w:space"), "8")
    el.set(qn("w:color"), colour)
    bdr.append(el)
    pPr.append(bdr)


def bottom_rule(par, colour=NAVY_HEX, sz="8"):
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    el = OxmlElement("w:bottom")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "6")
    el.set(qn("w:color"), colour)
    bdr.append(el)
    pPr.append(bdr)


def keep_next(par):
    pPr = par._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))


# ---------- content helpers ----------

def para(text="", size=10.5, bold=False, italic=False, colour=None,
         after=6, before=0, align=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if colour is not None:
            r.font.color.rgb = colour
    return p


def rich(segments, size=10.5, after=6, indent=None):
    """segments = [(text, bold?), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    for text, bold in segments:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
    return p


def h1(text):
    p = para(text, size=15, bold=True, colour=NAVY, before=16, after=4)
    bottom_rule(p, RULE, "6")
    keep_next(p)
    return p


def h2(text):
    p = para(text, size=11.5, bold=True, colour=NAVY, before=9, after=3)
    keep_next(p)
    return p


def bullets(items, numbered=False):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(5)
        marker = f"{i}.  " if numbered else "•  "
        r = p.add_run(marker)
        r.font.size = Pt(10.5)
        r.bold = numbered
        if numbered:
            r.font.color.rgb = NAVY
        r2 = p.add_run(t)
        r2.font.size = Pt(10.5)


def caption(text):
    para(text, size=8.5, italic=True, colour=GREY, after=10)


def table(header, rows, widths, first_col_bold=False):
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t._tbl.tblPr.append(_fixed_layout())
    set_borders(t)

    total = sum(widths)
    inches = [6.5 * w / total for w in widths]

    hdr = t.rows[0].cells
    for i, txt in enumerate(header):
        hdr[i].width = Inches(inches[i])
        shade(hdr[i], HDRFILL)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = NAVY

    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].width = Inches(inches[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(str(txt))
            r.font.size = Pt(8.5)
            if first_col_bold and i == 0:
                r.bold = True
    para("", after=2)
    return t


def figure(fname, width_in=4.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(os.path.join(CH, fname), width=Inches(width_in))


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def footer_text():
    f = doc.sections[0].footer
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Edwards  ·  Colorado Medicare Part D Prescribing Pattern Analysis  ·  page ")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    # PAGE field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rPr.append(sz)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "4A5568"); rPr.append(col)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = "1"; run.append(t)
    fld.append(run)
    p._p.append(fld)


# =====================================================================
# TITLE BLOCK
# =====================================================================
para("Finding the Prescribers a Spend Ranking Cannot See",
     size=21, bold=True, colour=NAVY, before=26, after=2)
p = para("A Colorado Medicare Part D Prescribing Pattern Analysis, 2024",
         size=13, colour=GREY, after=10)
bottom_rule(p, NAVY_HEX, "8")
para("Caleb Edwards, M.D.", size=11, bold=True, after=1)
para("Data Analytics Capstone  ·  Front Range Community College / Upright Education",
     size=10, colour=GREY, after=1)
para("August 2026", size=10, colour=GREY, after=14)

rich([("Summary.  ", True),
      ("Colorado's 2024 Medicare Part D file records $2.74 billion in drug cost across "
       "19,390 prescribers. A population health team can review a few hundred of them. "
       "This analysis shows that the review list most organizations build — a statewide "
       "spend or unit-cost ranking — misses roughly half of the prescribers who are "
       "genuinely unusual for their own specialty, and that comparing each prescriber "
       "to their own peer group instead yields a tractable, defensible cohort of 507 "
       "clinicians holding $317.6 million in above-peer spend.", False)],
     size=10.5, after=12)

# =====================================================================
# 1. BUSINESS QUESTION
# =====================================================================
h1("1. The business question")

para("Medicare Part D is the outpatient prescription drug benefit for people on Medicare, "
     "delivered through private plans approved by CMS. Every year CMS publishes a public "
     "file recording what each prescriber in the country prescribed and what those "
     "prescriptions cost. The 2024 Colorado extract covers $2.74 billion in drug cost "
     "across 19,390 prescribers and 1,177 generic drug molecules. It contains no patient "
     "identifiers and no protected health information.")

para("A population health team inside a health system or a plan has a handful of "
     "pharmacists and perhaps one medical director. It cannot review nineteen thousand "
     "prescribers. It can review a few hundred. The decision this project supports is "
     "therefore narrow and practical:")

p = para("Which prescriber segments should a health system prioritize for outreach, "
         "education, or further review — and on what evidence?",
         size=11, bold=True, colour=NAVY, indent=0.3, after=9)
left_bar(p)

para("There are two ways to build that review list, and the choice between them is the "
     "analytical heart of this project.")

rich([("Option A — rank everyone statewide.  ", True),
      ("Sort all 19,390 prescribers by total spend, or by cost per prescription, and take "
       "the top of the list. It is simple, fast, and it is what most programs do. Its "
       "weakness is that an oncologist and a dentist end up on the same ranking. Cancer "
       "drugs cost thousands of dollars per prescription because they are cancer drugs. A "
       "statewide ranking largely measures what kind of medicine a clinician practices, "
       "not how they practice it.", False)])

rich([("Option B — compare each prescriber to their own specialty.  ", True),
      ("Rank each prescriber only against peers in the same specialty, then flag the ones "
       "sitting above their own peer group. It requires the full file and more work to "
       "build, but it holds case mix roughly constant: family practice is compared to "
       "family practice, oncology to oncology.", False)])

para("My working hypothesis was that the harder list is worth building. This report tests "
     "that hypothesis and, having confirmed it, converts it into a specific cohort a team "
     "could begin reviewing on Monday morning. The application is direct: peer-comparison "
     "outreach is an established tool in pharmacy utilization management, and the "
     "constraint it always runs into is reviewer capacity. Getting the targeting right is "
     "what determines whether the program pays for itself.")

# =====================================================================
# 2. DATA
# =====================================================================
h1("2. Data collection and preparation")

rich([("Source.  ", True),
      ("CMS Medicare Part D Prescribers by Provider and Drug, 2024, Colorado extract, "
       "downloaded from data.cms.gov. The file has 390,473 rows and 22 columns. Each row "
       "is one prescriber, one brand name, one generic name, for the calendar year — "
       "carrying total claims, 30-day fills, day supply, drug cost, and beneficiary "
       "counts, plus a 65-and-older subgroup breakout. It is a public use file with no PHI.",
       False)])

para("I read the raw CSV exactly once for the entire project and cached a single cleaned "
     "copy that every later stage reads. Preparation was deliberately conservative. Ten of "
     "the columns are numeric by definition but load as text, so I coerced them to numbers; "
     "not one non-blank value failed to parse, and the sum of drug cost was identical "
     "before and after coercion at $2,737,455,388.61. I checked for duplicate rows across "
     "all 22 columns and on the file's grain key of NPI plus brand plus generic, and found "
     "none. I checked whether any NPI mapped to more than one specialty or more than one "
     "city, and none did — which means specialty and city can be carried from row level to "
     "prescriber level without loss.")

rich([("The judgment call worth naming is that I did not fill in blanks.  ", True),
      ("In this file a blank is not an omission — CMS suppresses any cell that would "
       "identify a small number of patients. Filling those with zero would have invented "
       "data. Every blank was left blank, and no row was filtered, dropped, or reweighted "
       "at any step. 390,473 rows in, 390,473 rows out.", False)])

table(
    ["Measure", "Raw CSV", "Cleaned file", "Delta"],
    [
        ["Rows", "390,473", "390,473", "0"],
        ["Columns", "22", "24 (22 source + 2 derived)", "+2"],
        ["SUM(Tot_Drug_Cst)", "$2,737,455,388.61", "$2,737,455,388.61", "$0.00"],
        ["SUM(Tot_Clms)", "16,573,710", "16,573,710", "0"],
        ["Distinct prescribers (NPI)", "19,390", "19,390", "0"],
        ["Distinct specialties", "97", "97", "0"],
        ["Distinct generic molecules", "1,177", "1,177", "0"],
    ],
    [32, 22, 32, 14],
)
caption("Reconciliation between the raw extract and the cached clean file. "
        "The step-by-step record, with a row count before and after each of ten steps, "
        "is in CLEANING.md.")

h2("The choice of unit metric")

para("One preparation decision shaped everything downstream, so it deserves its own "
     "treatment. The measure a clinician would naturally reach for is cost per patient. I "
     "could not compute it. The total beneficiary column is blank on 58.9% of rows under "
     "CMS privacy suppression, and the 65-and-older beneficiary column is blank on 86.3%. "
     "A denominator that is missing more than half the time is not a denominator, so no "
     "per-patient rate appears anywhere in this analysis.")

para("What is complete is claims and dollars — both 0% blank. The unit metric throughout "
     "is therefore cost per claim, computed as total drug cost divided by total claims, "
     "with cost per 30-day fill as a companion measure. Both denominators are guarded "
     "against zero, and because CMS suppresses any prescriber-drug combination below 11 "
     "claims, neither guard ever fired.")

para("That distribution is severely right-skewed: the median cost per claim is $31.09 "
     "against a mean of $193.65. Skew of that severity rules out treating a test that "
     "assumes normality as the primary evidence, which is why the statistics in Section 3 "
     "lead with a rank-based test.")

figure("stage3_01_cost_per_claim_distribution.png", 4.5)
caption("Figure 1 — Distribution of cost per claim. The long right tail is what drives the "
        "choice of a rank-based test. Source: capstone_analysis.ipynb, chart stage3_01.")

# =====================================================================
# 3. METHODOLOGY
# =====================================================================
h1("3. Methodology")

para("The pipeline is four tools in sequence, each layer checking the one before it. This "
     "structure follows the proposal, and I kept to it.")

bullets([
    "Excel. Aggregate tables built straight off the raw extract — the fast first pass "
    "across prescriber, specialty, and drug cuts, and the tie-out that every later number "
    "is checked against. All sheets reconcile to $2,737,455,388.61.",
    "SQL. The full 390,473-row file loaded to SQLite as the system of record. Percentile "
    "window functions — NTILE() partitioned by specialty — build the peer bands and "
    "identify outliers inside each specialty. Every query carries a plain-language comment "
    "stating what it does and why.",
    "Python. pandas and seaborn. This layer tests which dimension actually explains cost, "
    "ranks each prescriber inside their own specialty, and produces the seven charts. The "
    "notebook runs clean-kernel end to end: 23 code cells, zero errors.",
    "Tableau. The executive layer — the prioritized segments in a form a population health "
    "director can act on — paired with a one-page recommendation memo.",
], numbered=True)

h2("Peer groups and the outlier rule")

para("A peer group is a prescriber specialty. To keep the comparison honest I restricted "
     "within-specialty work to specialties with at least 30 prescribers; a specialty with "
     "four members produces a median nobody should trust. That rule qualifies 46 of 97 "
     "specialties, covering 18,940 of 19,390 prescribers. The 51 excluded specialties "
     "account for 2.3% of prescribers and 3.0% of spend, and I claim nothing about them.")

para("A prescriber is flagged as a within-specialty outlier when their cost per claim "
     "exceeds the 95th percentile of their own specialty. The parallel statewide rule, "
     "used only for comparison, flags the top 5% of all prescribers on cost per claim "
     "regardless of specialty. The SQL layer builds the same construct with NTILE(20) band "
     "20; the two rules agree closely but not exactly, and where both are quoted the "
     "notebook rule governs.")

h2("A substitution I want to state plainly")

rich([("My proposal called for regression to identify which dimensions drive cost. "
       "Regression beyond a fitted line was never covered in this curriculum — no "
       "coefficient tables, no R², no statsmodels or scikit-learn. Rather than import a "
       "method I could not defend line by line, I answered the same question with "
       "techniques that were taught: ", False),
      ("one-way ANOVA and the Kruskal–Wallis test", True),
      (" to compare cost per claim across specialty and across city, and ", False),
      ("within-specialty percentile ranking", True),
      (" using groupby, quantile, and merge to locate individual prescribers against their "
       "peers. Two smaller substitutions follow the same principle: a Pareto "
       "cumulative-percent line was replaced by a bar chart with the concentration figure "
       "stated in the title, and log-scale axes were replaced by 99th-percentile trims "
       "declared on the chart rather than hidden.", False)])

para("ANOVA and Kruskal–Wallis tell you whether a grouping explains variation; the "
     "percentile ranking tells you which individuals sit outside their group. Together, "
     "that is the entire business question, and I can defend every line of it.")

# =====================================================================
# 4. FINDINGS
# =====================================================================
h1("4. Key findings")

h2("Finding 1 — Specialty drives cost. Geography barely matters.")

para("I tested two ways of grouping prescribers, specialty and city, to see which explains "
     "cost per claim. Both came back statistically significant, and that is not by itself "
     "a result — with roughly 19,000 prescribers almost anything rejects the null. The "
     "size of the effect is the finding.")

table(
    ["Dimension", "Qualifying groups", "Kruskal–Wallis H", "ANOVA F",
     "Span of group medians", "p75 / p25 fold"],
    [
        ["Specialty", "46 of 97", "8,608.94", "51.83", "$3.89 → $2,272.07  (584×)", "11.23×"],
        ["City", "60 of 226", "254.96", "3.17", "$4.38 → $79.71  (18×)", "1.81×"],
    ],
    [15, 15, 16, 11, 27, 16],
    first_col_bold=True,
)
caption("Both p-values fall below 0.001; specialty's is below 1e-300. Kruskal–Wallis is the "
        "test relied on because cost per claim is severely right-skewed; ANOVA is reported "
        "alongside it and points the same way. Source: FINDINGS.md F1.")

para("The middle half of specialties differ eleven-fold on median cost per claim. The "
     "middle half of cities differ 1.8-fold, which is close to noise. Across the full "
     "range, the cheapest qualifying specialty sits at $3.89 per claim and the most "
     "expensive at $2,272.07 — a 584-fold spread that is case mix, not behavior.")

rich([("The practical consequence is the one that matters: ", False),
      ("cost-per-claim comparisons must be made inside a specialty.", True),
      (" Comparing an oncologist to a dentist measures case mix. Comparing a Denver "
       "prescriber to a Pueblo prescriber measures almost nothing. Geography is off the "
       "table as a targeting dimension.", False)])

h2("Finding 2 — Spending a lot and charging a lot are unrelated.")

para("I checked whether the prescribers who spend the most are also the ones whose "
     "individual prescriptions are most expensive. The Pearson correlation between cost "
     "per claim and total claims is −0.011. That is zero.")

table(
    ["Prescriber-level pair", "Pearson r"],
    [
        ["Total cost vs total claims", "0.380"],
        ["Cost per claim vs total claims", "−0.011"],
        ["Cost per claim vs total cost", "0.337"],
    ],
    [70, 30],
)
caption("Correlations computed at prescriber level on the qualifying population. "
        "Source: FINDINGS.md F4, chart stage3_07.")

para("Read plainly: knowing that a prescriber spent two million dollars tells you nothing "
     "about whether their prescriptions are expensive. High spend usually means high "
     "volume — a busy clinic seeing a lot of patients. This is the analytic justification "
     "for the whole project. If the two measures moved together a spend ranking would be an "
     "acceptable shortcut; because they do not, a high-spend review list and a "
     "high-unit-cost review list are lists of different people.")

h2("Finding 3 — Half the prescribers who stand out among peers are invisible statewide.")

para("This is the direct test of the fork described in Section 1. I built both lists and "
     "compared them.")

table(
    ["Flagged by", "Prescribers", "Reading"],
    [
        ["Statewide top 5% only", "431",
         "High unit cost, but normal for their specialty — explained by case mix"],
        ["Both methods", "498",
         "High unit cost and unusual against their own peers"],
        ["Within-specialty top 5% only", "467",
         "Unusual against their own peers, invisible on a statewide list"],
    ],
    [26, 13, 61],
    first_col_bold=True,
)
caption("Flag rule: cost per claim above the prescriber's own specialty 95th percentile, "
        "versus above the statewide 95th percentile. Source: FINDINGS.md F2, chart stage3_06.")

para("Of the 965 prescribers sitting above their own specialty's 95th percentile, 467 — "
     "48% — never appear on a statewide top-5% list at all. Conversely, 431 of the 929 "
     "statewide flags are fully accounted for by specialty. A statewide ranking is roughly "
     "half noise for this purpose, and it misses about half the real signal.")

para("Those 965 prescribers are 5.1% of the state and carry 22.1% of qualifying spend. "
     "Their median total drug cost is $144,422 against $4,924 for everyone else — a "
     "29-fold gap on the median, significant at H = 658.63, p = 3.0e-145. I deliberately "
     "ran that test on total cost rather than on cost per claim, because cost per claim "
     "was the variable used to build the flag and re-testing it would have been circular.")

figure("stage3_06_within_specialty_outliers_scatter.png", 4.4)
caption("Figure 2 — Prescribers by claim volume and total cost, with within-specialty "
        "outliers highlighted. Flagged prescribers sit high on total cost at low claim "
        "counts. Source: chart stage3_06.")

h2("Finding 4 — The money is concentrated, which is what makes a targeted program viable.")

table(
    ["Cut", "Share of the $2.74B"],
    [
        ["Top 10 of 1,177 generic molecules", "33.5%  ($917.8M)"],
        ["Top 20 generic molecules", "45.5%  ($1.25B)"],
        ["Top 100 generic molecules", "79.0%"],
        ["Top 100 of 19,390 prescribers", "18.5%"],
        ["Top 1,000 prescribers", "56.1%"],
        ["Top 10% of prescribers (1,939)", "71.9%  ($1.97B)"],
    ],
    [62, 38],
)
caption("Concentration on the drug axis and the prescriber axis. Source: FINDINGS.md F5, "
        "chart stage3_04.")

para("Twenty molecules — 1.7% of the drug list — carry 45.5% of Colorado Part D spend, and "
     "the top decile of prescribers carries 71.9%. This is a feasibility finding rather "
     "than a fairness one: because the money is concentrated rather than spread evenly, a "
     "formulary or utilization review scoped to twenty molecules touches nearly half the "
     "state's spend.")

# =====================================================================
# 5. RECOMMENDATIONS
# =====================================================================
h1("5. Recommendations and next steps")

para("The table below ranks specialties by the dollars sitting above that specialty's own "
     "95th percentile — the money a peer-based review would actually be looking at. It is "
     "the notebook's output sorted, not a judgment layered on top of it.")

table(
    ["#", "Specialty", "Outliers / prescribers", "Dollars above own p95",
     "% of specialty spend", "Cost per claim"],
    [
        ["1", "Nurse Practitioner", "176 / 3,509", "$109,343,136", "27.8%", "$137"],
        ["2", "Physician Assistant", "151 / 3,010", "$106,670,500", "41.6%", "$147"],
        ["3", "Internal Medicine", "65 / 1,292", "$60,293,209", "19.1%", "$98"],
        ["4", "Family Practice", "115 / 2,288", "$41,334,113", "11.7%", "$73"],
        ["5", "Hematology-Oncology", "6 / 115", "$33,045,826", "14.5%", "$2,652"],
        ["6", "Urology", "9 / 178", "$26,685,838", "53.0%", "$320"],
        ["7", "Pulmonary Disease", "8 / 154", "$25,642,076", "24.8%", "$907"],
        ["8", "Cardiology", "11 / 205", "$24,133,908", "17.9%", "$272"],
    ],
    [5, 23, 20, 21, 18, 13],
)
caption("Priority segments ranked by outlier dollars. Rows 1–4 are the recommended first "
        "cohort. Source: FINDINGS.md, “Recommended priority segments”.")

h2("Primary recommendation — review the top four segments first, as a single cohort of 507 clinicians.")

para("Nurse Practitioner, Physician Assistant, Internal Medicine and Family Practice are "
     "the four cheapest segments on this list, at $73 to $147 per claim, and yet their "
     "within-specialty outliers hold $317.6 million — 54% of all outlier dollars in "
     "Colorado. That combination is the entire argument for using a within-specialty rule.")

para("These 507 prescribers are invisible to any statewide unit-cost screen: a $600 claim "
     "looks unremarkable next to oncology and extraordinary next to a family practice peer "
     "group whose median is $73. And 507 people is a tractable review list — a project a "
     "small team can finish, not a program it has to staff.")

h2("Second tier — Urology, on concentration rather than size.")

para("Nine prescribers hold 53.0% of the specialty's entire spend. No other segment on the "
     "list is close to that concentration. Nine charts is roughly a day of work, and the "
     "potential return per reviewer hour is the highest on the table.")

h2("Deprioritize the expensive specialties relative to their dollar rank.")

para("Hematology-Oncology, Pulmonary Disease and Cardiology appear high in the table "
     "because their drugs are genuinely expensive — $2,652, $907 and $272 per claim. "
     "Finding 1 establishes that specialty explains most of that. The residual "
     "within-specialty signal is real but small, and the clinical justification for "
     "high-cost oncology and pulmonary agents is correspondingly stronger. Reviewer hours "
     "are better spent where the peer gap is not already explained by the pharmacology.")

h2("Next steps")

bullets([
    "Run the 507-clinician cohort as a pilot. Peer-comparison letters and pharmacist "
    "outreach, then measure in the next data year whether the peer gap narrowed. That is a "
    "measurable intervention rather than a report that sits on a shelf.",
    "Scope a formulary review to the top twenty molecules. Those twenty carry 45.5% of "
    "state spend, which makes them the cheapest structural lever available.",
    "Rebuild the flag with clinical context. Diagnoses and panel composition would turn a "
    "statistical outlier into an actionable one. That data lives inside a health system, "
    "not in a public file.",
    "Repeat the analysis on the 2023 file. One year cannot separate a persistent pattern "
    "from a single unusual year. Two can.",
], numbered=True)

# =====================================================================
# 6. LIMITATIONS
# =====================================================================
h1("6. Limitations")

para("This is data about real clinicians, so I want to be precise about what it does not say.")

for label, text in [
    ("No clinical context.",
     "The file carries no diagnoses, no severity, and no patient panel composition. A "
     "prescriber above their peer group is a statistical outlier — a reason to look at a "
     "chart. It is not a finding of waste, and it is not an accusation."),
    ("Cost is plan-paid, not net.",
     "Tot_Drug_Cst is what CMS reports as paid. It is not net of rebates or discounts, so "
     "cross-drug comparisons overstate the spread on heavily rebated brands."),
    ("No cost per beneficiary.",
     "Beneficiary counts are suppressed on 58.9% of rows, so no per-patient rate is "
     "computed anywhere in this analysis. Cost per claim and cost per 30-day fill are the "
     "only unit metrics used."),
    ("CMS suppression sets a floor on the file.",
     "Every published row has at least 11 claims. A prescriber's totals here are "
     "therefore sums over only the rows CMS chose to report, which understates the true "
     "total for anyone whose volume is spread thinly across many drugs."),
    ("Specialty labels are not uniformly sourced.",
     "Prscrbr_Type is derived from claims for 88.2% of rows and from the NPPES registry "
     "for the rest, so a peer group built on this column mixes two derivations."),
    ("Scope, and what remains open.",
     "2024 Colorado only — no trend and no national benchmark. The ≥30-prescriber rule "
     "excludes 51 of 97 specialties, covering 2.3% of prescribers and 3.0% of spend. "
     "Fifteen open data questions remain logged in QUESTIONS.md — none large enough to "
     "move the figures in this report, and none silently fixed."),
]:
    rich([(label + "  ", True), (text, False)], after=6)

# =====================================================================
# APPENDIX
# =====================================================================
page_break()
h1("Appendix — Artifacts")

para("Every file in the outputs folder and what it demonstrates. Items marked “working "
     "file” document the process and were not written as graded prose.")

h2("Deliverables")
table(
    ["File", "What it demonstrates"],
    [
        ["Edwards_Caleb_Capstone_Report.docx",
         "This report — business question, data collection and preparation, methodology, "
         "findings, recommendations, limitations, and this appendix."],
        ["Edwards_Caleb_Capstone_Deck.pptx",
         "Fourteen slides with full speaker notes, written for an audience with no prior "
         "knowledge of Medicare Part D."],
        ["RECOMMENDATION_MEMO.docx",
         "One-page memo to a population health leadership team. All 34 figures in it were "
         "checked programmatically against FINDINGS.md with zero mismatches."],
        ["Capstone_Dashboard.twbx",
         "Tableau executive dashboard — outlier dollars by segment, cost per claim by "
         "segment, and a volume-versus-unit-cost scatter across the top 8 priority segments."],
        ["tableau_prescriber_segments.csv",
         "The prescriber-level extract feeding the dashboard; reconciles exactly to the "
         "priority table in Section 5."],
    ],
    [34, 66],
)

h2("Analysis and code")
table(
    ["File", "What it demonstrates"],
    [
        ["capstone_analysis.ipynb",
         "The Python analysis. 23 cells, clean-kernel run, zero errors — ANOVA, "
         "Kruskal–Wallis, within-specialty percentile ranking, correlations, and all seven "
         "charts. Every cell carries a plain-language comment."],
        ["capstone_segmentation.sql",
         "SQL segmentation using CTEs and NTILE() window functions to build peer bands and "
         "identify outliers inside each specialty. Commented query by query."],
        ["part_d_profiling.sql",
         "Profiling SQL — 35 queries establishing the shape, completeness, and "
         "distribution of the raw file."],
        ["CAPSTONE_EXCEL_AGGREGATES.xlsx",
         "Excel aggregate tables built from the raw extract, tying to $2,737,455,388.61 — "
         "the reference every later number is checked against."],
        ["charts/  (7 PNG files)",
         "The complete chart set: cost-per-claim distribution, cost by specialty (box and "
         "bar), top generic molecules by spend, group-median dispersion by dimension, "
         "within-specialty outlier scatter, and the measure correlation heatmap."],
        ["part_d_co_clean.csv",
         "The single cached clean file. The raw CSV is read once; every later stage reads "
         "this instead."],
        ["part_d.sqlite",
         "The full 390,473-row file as a SQLite database — the SQL system of record."],
        ["Build scripts (load_part_d.py, build_excel_aggregates.py, build_nb.py, "
         "make_charts.py, run_segmentation_sql.py, build_tableau_extract.py, "
         "build_recommendation_memo.py, build_report.py)",
         "Every deliverable in this folder is reproducible from the raw CSV by running "
         "these in order."],
    ],
    [34, 66],
)

h2("Documentation")
table(
    ["File", "What it demonstrates"],
    [
        ["CLEANING.md",
         "The data cleaning log — ten steps, each with a row count before and after, plus a "
         "full data-limitations section. This is the preparation audit trail."],
        ["FINDINGS.md",
         "The findings of record. Every figure in this report, the deck, the memo and the "
         "dashboard traces to a numbered finding here."],
        ["QUESTIONS.md",
         "Fifteen open data questions raised and deliberately left undecided — ambiguities "
         "logged rather than silently resolved."],
        ["SQL_RESULTS.md, STATS_OUTPUT.md, PY_RESULTS.md",
         "Raw query and statistical output, retained so any number in this report can be "
         "traced back to the run that produced it."],
        ["PROGRESS.md",
         "Stage-by-stage run log, including decisions made without supervision and "
         "residual risks (working file)."],
        ["CALIBRATION.md, REQUIREMENTS.md, VERIFICATION.md",
         "Scope and technique controls — what the curriculum taught, what the rubric "
         "requires, and what was verified (working files)."],
        ["CHART_FLAGS.md, SQL_FLAGS.md, PY_FLAGS.md, RUN_LOG.md, PREGAME_RESEARCH.md",
         "Earlier exploratory work, retained for completeness (working files)."],
        ["charts/superseded/  (5 PNG files)",
         "Five earlier figures built with techniques outside the curriculum. Moved rather "
         "than deleted, and excluded from the deliverable set."],
        ["CAPSTONE_AGGREGATES.xlsx, CAPSTONE_PROFILE.xlsx, CITY_REFERENCE.xlsx, "
         "outliers.csv, part_d_analysis.ipynb, regression_output.csv",
         "Prior exploratory work from before the pipeline was rebuilt on taught "
         "techniques. Superseded by the files above and retained, not deleted, so the "
         "path from first pass to final analysis stays visible (working files)."],
    ],
    [34, 66],
)

para("Reproducibility: every number in this report can be regenerated from "
     "Medicare_Part_D_Prescribers_by_Provider_and_Drug_2024.csv by reading it with "
     "encoding='utf-8-sig' and running the build scripts listed above in order.",
     size=9, italic=True, colour=GREY, after=4)

footer_text()

dest = os.path.join(OUT, "Edwards_Caleb_Capstone_Report.docx")
doc.save(dest)
print("written:", dest, os.path.getsize(dest), "bytes")
