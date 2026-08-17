from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GREY = RGBColor(0x4A, 0x55, 0x68)

doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(s, m, Inches(0.62))

st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10)
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.06

def _fixed_layout():
    el = OxmlElement("w:tblLayout"); el.set(qn("w:type"), "fixed")
    return el

def shade(cell, hexfill):
    el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto"); el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)

def border(par, edge="bottom", sz=12, color="1A365D"):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "4"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)

def para(runs, after=5, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(space_before)
    for t, bold, color in runs:
        r = p.add_run(t); r.bold = bold; r.font.size = Pt(10)
        r.font.name = "Calibri"
        if color is not None: r.font.color.rgb = color
    return p

def head(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = NAVY; r.font.name = "Calibri"
    return p

# ---- title ----
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
r = p.add_run("MEMORANDUM"); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
rule = doc.add_paragraph(); rule.paragraph_format.space_after = Pt(8)
rule.add_run("").font.size = Pt(1); border(rule)

# ---- To/From/Date/Re ----
hdr = doc.add_table(rows=0, cols=2)
hdr.autofit = False
hdr._tbl.tblPr.append(_fixed_layout())
for label, value in [
    ("To", "Population Health Leadership Team"),
    ("From", "Caleb Edwards, M.D. — Data Analytics"),
    ("Date", "August 12, 2026"),
    ("Re", "Targeting pharmacy utilization review using within-specialty cost "
           "benchmarks — Colorado Medicare Part D, 2024"),
]:
    row = hdr.add_row()
    row.cells[0].width = Inches(0.62); row.cells[1].width = Inches(6.64)
    c0 = row.cells[0].paragraphs[0]; c0.paragraph_format.space_after = Pt(1)
    r0 = c0.add_run(label.upper()); r0.bold = True; r0.font.size = Pt(8.5); r0.font.color.rgb = GREY
    c1 = row.cells[1].paragraphs[0]; c1.paragraph_format.space_after = Pt(1)
    r1 = c1.add_run(value); r1.font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(0)

# ---- body ----
head("Recommendation")
para([("Stand up a within-specialty prescribing review and scope the first wave to 507 clinicians", True, None),
      (" across four segments — Nurse Practitioner, Physician Assistant, Internal Medicine, and "
       "Family Practice. Each sits above their own specialty's 95th percentile for cost per claim; "
       "together they account for $317.6M, 54% of all flagged Part D dollars in Colorado. Add Urology "
       "as a small second wave — nine prescribers hold 53.0% of that specialty's entire spend.", False, None)])

head("Why the usual screen misses them")
para([("A statewide high-cost list is a list of oncologists and geneticists — expensive because "
       "their drugs are expensive, not because their prescribing is unusual. Of the 965 Colorado "
       "prescribers who exceed their own specialty's 95th percentile, ", False, None),
      ("467 (48%) never appear on a statewide top-5% list at all", True, None),
      (". Conversely, 431 of 929 statewide flags are fully explained by case mix. Cost per claim is also "
       "effectively uncorrelated with claim volume (r = −0.011), so a spend-ranked review list and a "
       "unit-cost-ranked list find different people. Ranking spend alone will not surface this cohort.", False, None)])

head("What the flagged cohort represents")
para([("Across the 46 Colorado specialties with at least 30 prescribers, 965 clinicians — 5.1% — "
       "carry 22.1% of qualifying spend ($587.1M against $2.74B statewide). Their median total drug cost is "
       "29× the median of every other prescriber (Kruskal–Wallis H = 658.6, p < 0.001). That test runs "
       "on total cost, not the unit-cost measure used to build the flag, so the gap is not an artifact "
       "of the ranking.", False, None)])

# ---- table ----
head("First-wave priority segments")
data = [
    ("Segment", "Clinicians", "Spend flagged", "% of segment", "$/claim"),
    ("Nurse Practitioner", "176", "$109.3M", "27.8%", "$137"),
    ("Physician Assistant", "151", "$106.7M", "41.6%", "$147"),
    ("Internal Medicine", "65", "$60.3M", "19.1%", "$98"),
    ("Family Practice", "115", "$41.3M", "11.7%", "$73"),
    ("Tier 1 total", "507", "$317.6M", "—", "—"),
    ("Urology (Tier 2)", "9", "$26.7M", "53.0%", "$320"),
]
widths = [Inches(2.15), Inches(1.05), Inches(1.35), Inches(1.15), Inches(0.80)]
tbl = doc.add_table(rows=0, cols=5)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for ri, vals in enumerate(data):
    row = tbl.add_row()
    for ci, v in enumerate(vals):
        cell = row.cells[ci]; cell.width = widths[ci]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
        if ci > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(v); r.font.size = Pt(9); r.font.name = "Calibri"
        if ri == 0:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(cell, "1A365D")
        elif ri == 5:
            r.bold = True; shade(cell, "EDF2F7")

cap = doc.add_paragraph(); cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(4)
rc = cap.add_run("Spend flagged = 2024 Part D drug cost held by prescribers above their own specialty's "
                 "95th percentile cost per claim. Source: CMS Medicare Part D Prescribers by Provider "
                 "and Drug, 2024, Colorado.")
rc.italic = True; rc.font.size = Pt(7.5); rc.font.color.rgb = GREY

head("Suggested next steps")
for t in [
    "1.  Pull the 507 NPIs and cross-reference against your attributed panels; expect only a subset to be in-network.",
    "2.  Scope the drug review narrowly. The top 20 generic molecules carry 45.5% of Colorado Part D spend, so a formulary and utilization review touching 20 molecules covers nearly half the money.",
    "3.  Open with education, not audit. Share each clinician's cost per claim against their own specialty's distribution — a peer benchmark most prescribers have never seen.",
    "4.  Deprioritize Hematology-Oncology, Pulmonary Disease, and Cardiology relative to their raw dollar rank — their high unit costs are largely explained by the specialty itself, and the clinical justification for those agents is stronger.",
]:
    para([(t, False, None)], after=3)

foot_rule = doc.add_paragraph(); foot_rule.paragraph_format.space_before = Pt(8)
foot_rule.paragraph_format.space_after = Pt(4)
foot_rule.add_run("").font.size = Pt(1); border(foot_rule, "top", 6, "CBD5E0")

para([("Important limitation. ", True, GREY),
      ("This dataset carries no diagnoses, no severity, and no panel composition. A within-specialty "
       "outlier is a statistical outlier — a reason to look, not a finding of waste or impropriety. "
       "Reported cost is plan-paid, not net of rebates, and beneficiary counts are suppressed on 58.9% of "
       "rows, so no per-patient rate is computed anywhere in this analysis.", False, GREY)], after=0)

doc.save("/tmp/memo/RECOMMENDATION_MEMO.docx")
print("saved")
